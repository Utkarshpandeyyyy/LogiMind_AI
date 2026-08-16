import os
import sys
import subprocess

def install_and_import():
    try:
        import docx
    except ImportError:
        print("python-docx not found. Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
    return docx

docx = install_and_import()
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Define color palette
color_primary = RGBColor(15, 23, 42) # Slate-900
color_secondary = RGBColor(71, 85, 105) # Slate-600

# Set title
title = doc.add_paragraph()
title_run = title.add_run("LogiMind AI - Presentation Slide Content")
title_run.bold = True
title_run.font.size = Pt(22)
title_run.font.color.rgb = color_primary
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add spacer
doc.add_paragraph("")

slides_data = [
    {
        "title": "Slide 1: Technology Stack Architecture",
        "subtitle": "High-performance, lightweight, and interactive decision intelligence architecture.",
        "sections": [
            ("Frontend (Presentation Layer):", [
                "Streamlit (v1.35+): Renders the dynamic executive dashboards, simulators, and copilot chat views.",
                "HTML5/Custom CSS: Embedded styling to create high-premium dark mode card designs and hero layouts."
            ]),
            ("Middleware & AI (Logic Layer):", [
                "LangGraph & LangChain: Powers the stateful, cyclic AI Copilot agent for intent classification and text-to-SQL logic.",
                "NetworkX (v3.2): Models supply chain nodes (hubs, customers) and routes as a topological graph network.",
                "Pandas & NumPy: Drives fast, vectorized in-memory data processing, delay simulations, and value-at-risk calculations."
            ]),
            ("Storage & Infrastructure (Data Layer):", [
                "PostgreSQL: Relational database storing live shipments, warehouse metadata, customer orders, and vehicle details.",
                "Docker Compose: Orchestrates PostgreSQL database and local development services containerization."
            ])
        ]
    },
    {
        "title": "Slide 2: Frontend & Interactive GIS Maps",
        "subtitle": "Real-time visualization and immersive user experience.",
        "sections": [
            ("Streamlit Framework:", [
                "Eliminates separate JS/HTML frameworks to provide pure Python hot-reloading dashboard UI.",
                "Uses session-state caching (st.cache_data) for fluid rendering of large shipment datasets."
            ]),
            ("Plotly Mapbox Integration:", [
                "Dynamically plots origins, destinations, and connecting freight lanes using Scattermapbox.",
                "Visualizes delay severity using color-coded nodes (Low, Medium, High, Critical) in dark-mode style."
            ]),
            ("Responsive Widgets:", [
                "Sliders, metrics cards, and selectboxes allow executive operators to run instant multi-factor simulations."
            ])
        ]
    },
    {
        "title": "Slide 3: Backend Database & Graph Network Topology",
        "subtitle": "Relational transactional data coupled with network dependency graphs.",
        "sections": [
            ("PostgreSQL Database:", [
                "Implements relational schema connecting shipments -> orders -> vehicles -> warehouses.",
                "Ensures data integrity and structured queries for precise supply chain audits."
            ]),
            ("NetworkX Topology Analysis:", [
                "Models physical logistics supply chain as a mathematical graph (nodes = hubs/warehouses, edges = shipping lanes).",
                "Calculates graph connectivity to identify downstream vulnerabilities when critical hubs (e.g., Chennai Hub) fail."
            ])
        ]
    },
    {
        "title": "Slide 4: Stateful AI Copilot (Natural Language Processing)",
        "subtitle": "Text-to-SQL AI Agent for conversational supply chain tracking.",
        "sections": [
            ("LangGraph Orchestration:", [
                "Builds stateful agent graph (AgentState) that classifies user questions (ETA requests, delay root causes, vehicle details).",
                "Bypasses static lookup tables with dynamic intent-driven routing."
            ]),
            ("Text-to-SQL Engine:", [
                "Converts conversational text into parametrized PostgreSQL queries (safeguarded from SQL injection).",
                "Executes against live Postgres databases and translates SQL results back into human-friendly explanations."
            ])
        ]
    }
]

for slide in slides_data:
    # Slide title
    h = doc.add_paragraph()
    r = h.add_run(slide["title"])
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = color_primary
    
    # Subtitle
    sub = doc.add_paragraph()
    sub_run = sub.add_run(slide["subtitle"])
    sub_run.italic = True
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = color_secondary
    
    # Content
    for category, points in slide["sections"]:
        p_cat = doc.add_paragraph()
        r_cat = p_cat.add_run(category)
        r_cat.bold = True
        r_cat.font.size = Pt(11.5)
        
        for pt in points:
            p_pt = doc.add_paragraph(style='List Bullet')
            r_pt = p_pt.add_run(pt)
            r_pt.font.size = Pt(10.5)
            
    doc.add_paragraph("—" * 40) # Divider line

output_dir = "documentation"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "PPT_CONTENT.docx")
doc.save(output_path)
print(f"\nSuccessfully generated DOCX at {os.path.abspath(output_path)}")
